import os
import docx
import pdfplumber
import subprocess

def extract_metadata(filepath):
    """
    Extracts metadata from the file (Author, Creator, etc.)
    Returns: Dict {"author": str, "last_modified_by": str, "creator": str, ...}
    """
    ext = os.path.splitext(filepath)[1].lower()
    metadata = {"author": "Unknown", "creator": "Unknown"}
    
    try:
        if ext == '.docx':
            doc = docx.Document(filepath)
            prop = doc.core_properties
            metadata["author"] = prop.author or "Unknown"
            metadata["last_modified_by"] = prop.last_modified_by or "Unknown"
            metadata["creator"] = "Microsoft Word (Implied)"
            
        elif ext == '.pdf':
            with pdfplumber.open(filepath) as pdf:
                # pdf.metadata is a dict
                raw_meta = pdf.metadata
                if raw_meta:
                    metadata["author"] = raw_meta.get("Author", "Unknown")
                    metadata["creator"] = raw_meta.get("Creator", "Unknown")
                    metadata["producer"] = raw_meta.get("Producer", "Unknown")
                    
        elif ext == '.doc':
            metadata["author"] = "Unknown (Legacy .doc)"
            
    except Exception as e:
        print(f"Error extracting metadata from {filepath}: {e}")
        
    return metadata

def extract_content(filepath):
    """
    Factory function to extract text with page metadata.
    Returns: List[{"text": str, "page": int}]
    """
    ext = os.path.splitext(filepath)[1].lower()
    if ext == '.docx':
        return extract_docx(filepath)
    elif ext == '.pdf':
        return extract_pdf(filepath)
    elif ext == '.doc':
        # legacy doc via antiword doesn't give page numbers easily, treat as single page
        return [{"text": extract_doc(filepath), "page": 1}]
    else:
        raise ValueError(f"不支持的文件格式: {ext} (仅支持 .docx, .doc, .pdf)")

def extract_text(filepath):
    """
    Legacy wrapper for string-only return (if used elsewhere).
    """
    content = extract_content(filepath)
    return "\n".join([item['text'] for item in content])

import shutil

def extract_docx(filepath):
    """
    Extracts text from a .docx file.
    STRATEGY: 
    1. Try to convert to PDF using LibreOffice (soffice).
    2. If successful, extract from PDF (preserves page numbers).
    3. If failed, fallback to python-docx (Page 1, no pagination).
    """
    if not os.path.exists(filepath):
        # This is the error user saw. Print detailed path for debugging.
        raise FileNotFoundError(f"文件不存在: {filepath}")

    import tempfile

    # Try converting to PDF first for pagination
    # Only if libreoffice is available
    soffice_bin = shutil.which('soffice') or shutil.which('libreoffice')
    print(f"DEBUG: Checking for LibreOffice... Found: {soffice_bin}")
    
    if soffice_bin:
        pdf_path = filepath.replace('.docx', '.pdf').replace('.DOCX', '.pdf')
        temp_home = None
        try:
            out_dir = os.path.dirname(filepath)
            
            # Create a dedicated temporary HOME directory for LibreOffice
            # This solves "User installation could not be completed" and permission errors
            temp_home = tempfile.mkdtemp()
            print(f"DEBUG: Created temp HOME for LibreOffice: {temp_home}")
            
            env = os.environ.copy()
            env['HOME'] = temp_home
            # Critical: Set XDG variables to prevent dconf from trying to write to /
            env['XDG_CACHE_HOME'] = os.path.join(temp_home, '.cache')
            env['XDG_CONFIG_HOME'] = os.path.join(temp_home, '.config')
            env['XDG_DATA_HOME'] = os.path.join(temp_home, '.local', 'share')
            env['XDG_RUNTIME_DIR'] = temp_home
                
            # Command: soffice --headless --convert-to pdf --outdir <dir> <file>
            cmd = [soffice_bin, '--headless', '--convert-to', 'pdf', '--outdir', out_dir, filepath]
            
            print(f"DEBUG: Executing command: {' '.join(cmd)}")
            
            # Run with timeout
            proc = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=60, env=env)
            
            print(f"DEBUG: LibreOffice Return Code: {proc.returncode}")
            if proc.returncode != 0:
                 print(f"DEBUG: LibreOffice Stderr: {proc.stderr.decode('utf-8', errors='ignore')}")
                 print(f"DEBUG: LibreOffice Stdout: {proc.stdout.decode('utf-8', errors='ignore')}")
            
            if proc.returncode == 0 and os.path.exists(pdf_path):
                print(f"DEBUG: PDF created successfully at {pdf_path}")
                # Success! Extract using PDF extractor logic
                try:
                    content = extract_pdf(pdf_path)
                    # Helper cleanup
                    try:
                        os.remove(pdf_path)
                    except:
                        pass
                    return content
                except Exception as e:
                    print(f"DEBUG: PDF Extraction failed after conversion: {e}")
                    # Change to fallback
            else:
                 print("DEBUG: PDF file not found after conversion (or return code non-zero)")
        except Exception as e:
            print(f"DEBUG: LibreOffice conversion failed with exception: {e}")
        finally:
            # Clean up temp HOME
            if temp_home and os.path.exists(temp_home):
                try:
                    shutil.rmtree(temp_home)
                    print(f"DEBUG: Removed temp HOME: {temp_home}")
                except Exception as cleanup_err:
                    print(f"DEBUG: Failed to remove temp HOME: {cleanup_err}")
    else:
        print("DEBUG: LibreOffice not found. Using XML fallback.")
    
    # Fallback to python-docx
    if not os.path.exists(filepath):
         raise FileNotFoundError(f"Fallback failed, file not found: {filepath}")

    doc = docx.Document(filepath)
    full_text = []
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text.strip())
    
    # Return as single 'page'
    return [{"text": "\n".join(full_text), "page": 1}]

def extract_doc(filepath):
    # antiword returns string
    # See previous implementation
    try:
        result = subprocess.run(
            ['antiword', filepath], 
            stdout=subprocess.PIPE, 
            stderr=subprocess.PIPE,
            text=True
        )
        if result.returncode != 0:
            if "No such file or directory" in str(result.stderr):
                 raise RuntimeError("系统未安装 antiword 工具。")
            raise RuntimeError(f"antiword 解析失败: {result.stderr.strip()}")
        return result.stdout.strip()
    except FileNotFoundError:
        if os.name == 'nt':
             raise RuntimeError("Windows 本地环境需手动安装 Antiword 或将文件转换为 .docx。\n服务器端(Docker)会自动安装支持。")
        raise RuntimeError("服务器未安装 antiword 工具，请运行 apt-get install antiword")
    except Exception as e:
        raise RuntimeError(f"解析 .doc 文件出错: {str(e)}")

def extract_pdf(filepath):
    """
    Extracts text from a .pdf file using pdfplumber.
    Returns: List[{"text": str, "page": int}]
    """
    content = []
    with pdfplumber.open(filepath) as pdf:
        for i, page in enumerate(pdf.pages):
            page_text = page.extract_text()
            if page_text:
                content.append({
                    "text": page_text, # Keep raw page text (with \n)
                    "page": i + 1
                })
    return content
